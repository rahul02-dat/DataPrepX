import matplotlib
matplotlib.use('Agg')

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from pathlib import Path
from typing import Dict, Any, List
from datetime import datetime
import math
import logging

from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle, PageBreak
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from PIL import Image as PILImage, ImageFile
DecompressionBombError = getattr(PILImage, "DecompressionBombError", Exception)

PIL_MAX_PIXELS = 2_000_000_000
PILImage.MAX_IMAGE_PIXELS = PIL_MAX_PIXELS
ImageFile.LOAD_TRUNCATED_IMAGES = True

from modules.utils import setup_logging, get_timestamp

logger = setup_logging()
log = logging.getLogger(__name__)

sns.set_style('whitegrid')
plt.rcParams['figure.figsize'] = (10, 6)
plt.ioff()


def safe_load_image(image_path: Path, width: float = None, height: float = None) -> RLImage:
    image_path = Path(image_path)
    
    old_limit = PILImage.MAX_IMAGE_PIXELS
    PILImage.MAX_IMAGE_PIXELS = None
    
    try:
        with PILImage.open(str(image_path)) as img:
            img.verify()
            img_width, img_height = img.size
    except Exception as e:
        log.error(f"Failed to verify image {image_path}: {e}")
        raise
    finally:
        PILImage.MAX_IMAGE_PIXELS = old_limit
    
    rl_image = RLImage(str(image_path))
    
    if width is not None or height is not None:
        if width is not None:
            width_pt = width * 72
        else:
            width_pt = None
        
        if height is not None:
            height_pt = height * 72
        else:
            height_pt = None
        
        if width_pt is not None and height_pt is not None:
            rl_image.drawWidth = width_pt
            rl_image.drawHeight = height_pt
        elif width_pt is not None:
            rl_image.drawWidth = width_pt
            aspect_ratio = rl_image.drawHeight / rl_image.drawWidth if rl_image.drawWidth else 1
            rl_image.drawHeight = width_pt * aspect_ratio
        elif height_pt is not None:
            rl_image.drawHeight = height_pt
            aspect_ratio = rl_image.drawWidth / rl_image.drawHeight if rl_image.drawHeight else 1
            rl_image.drawWidth = height_pt * aspect_ratio
    
    return rl_image


def safe_savefig(fig, path: Path, desired_dpi: int = 300, min_dpi: int = 50, max_width_inch: int = 18, max_height_inch: int = 12, **savefig_kwargs):
    path = Path(path)
    width_in, height_in = fig.get_size_inches()
    width_in = min(width_in, max_width_inch)
    height_in = min(height_in, max_height_inch)
    fig.set_size_inches(width_in, height_in)

    area_in = width_in * height_in

    try:
        max_dpi_allowed = int(math.floor(math.sqrt(PIL_MAX_PIXELS / area_in)))
    except Exception:
        max_dpi_allowed = 100

    dpi_to_use = desired_dpi
    if max_dpi_allowed < dpi_to_use:
        dpi_to_use = max(max_dpi_allowed, min_dpi)

    def pixels_for(dpi_val):
        return int(math.ceil(width_in * dpi_val)), int(math.ceil(height_in * dpi_val))

    w_px, h_px = pixels_for(dpi_to_use)
    if (w_px * h_px) > PIL_MAX_PIXELS:
        scale = math.sqrt(PIL_MAX_PIXELS / (width_in * height_in)) / dpi_to_use
        if scale <= 0 or scale > 1:
            scale = min(1.0, math.sqrt(PIL_MAX_PIXELS / (width_in * height_in)) / max(1, dpi_to_use))
        new_w = max(1.0, width_in * scale)
        new_h = max(1.0, height_in * scale)
        log.warning(f"safe_savefig: downscaling figure size from {(width_in, height_in)} to {(new_w, new_h)} to avoid Pillow limit")
        fig.set_size_inches(new_w, new_h)

        try:
            max_dpi_allowed = int(math.floor(math.sqrt(PIL_MAX_PIXELS / (new_w * new_h))))
        except Exception:
            max_dpi_allowed = min_dpi
        dpi_to_use = min(desired_dpi, max(max_dpi_allowed, min_dpi))

    try:
        fig.savefig(str(path), dpi=dpi_to_use, bbox_inches='tight', **savefig_kwargs)
    except DecompressionBombError:
        log.warning("Pillow DecompressionBombError on save; trying lower dpi and smaller size")
        for attempt in range(4):
            dpi_to_use = max(min_dpi, dpi_to_use // 2)
            try:
                fig.savefig(str(path), dpi=dpi_to_use, bbox_inches='tight', **savefig_kwargs)
                break
            except DecompressionBombError:
                continue
        else:
            log.error("safe_savefig: failed to save within safe limits; temporarily disabling Pillow limit to save file (risky).")
            old = getattr(PILImage, 'MAX_IMAGE_PIXELS', None)
            PILImage.MAX_IMAGE_PIXELS = None
            try:
                fig.savefig(str(path), dpi=min(desired_dpi, 150), bbox_inches='tight', **savefig_kwargs)
            finally:
                if old is not None:
                    PILImage.MAX_IMAGE_PIXELS = old


class ReportGenerator:
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.charts_dir = None

    def generate(self, df: pd.DataFrame, metadata: Dict[str, Any],
                 results: Dict[str, Any], output_dir: Path,
                 report_format: str = 'pdf') -> List[Path]:

        self.charts_dir = output_dir / 'charts'
        self.charts_dir.mkdir(parents=True, exist_ok=True)

        logger.info("Generating visualizations...")
        chart_paths = self._generate_charts(df, metadata, results)

        report_paths = []
        timestamp = get_timestamp()

        if report_format in ['pdf', 'both']:
            pdf_path = output_dir / f'DataPrepX_Report_{timestamp}.pdf'
            self._generate_pdf_report(df, metadata, results, chart_paths, pdf_path)
            report_paths.append(pdf_path)
            logger.info(f"PDF report generated: {pdf_path}")

        if report_format in ['docx', 'both']:
            docx_path = output_dir / f'DataPrepX_Report_{timestamp}.docx'
            self._generate_docx_report(df, metadata, results, chart_paths, docx_path)
            report_paths.append(docx_path)
            logger.info(f"DOCX report generated: {docx_path}")

        return report_paths

    def _generate_charts(self, df: pd.DataFrame, metadata: Dict[str, Any],
                         results: Dict[str, Any]) -> Dict[str, Path]:
        chart_paths = {}

        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()

        if numeric_cols:
            fig, ax = plt.subplots(figsize=(12, 8))
            df[numeric_cols[:10]].boxplot(ax=ax, rot=45)
            ax.set_title('Feature Distributions', fontsize=14, fontweight='bold')
            ax.set_ylabel('Value')
            plt.tight_layout()
            path = self.charts_dir / 'boxplot.png'
            safe_savefig(fig, path, desired_dpi=150, min_dpi=72)
            plt.close(fig)
            chart_paths['boxplot'] = path

        if len(numeric_cols) >= 2:
            corr_matrix = df[numeric_cols].corr()
            fig, ax = plt.subplots(figsize=(12, 10))
            sns.heatmap(corr_matrix, annot=True, fmt='.2f', cmap='coolwarm',
                        center=0, ax=ax, cbar_kws={'label': 'Correlation'})
            ax.set_title('Feature Correlation Matrix', fontsize=14, fontweight='bold')
            plt.tight_layout()
            path = self.charts_dir / 'correlation.png'
            safe_savefig(fig, path, desired_dpi=150, min_dpi=72)
            plt.close(fig)
            chart_paths['correlation'] = path

        if numeric_cols:
            fig, axes = plt.subplots(2, 2, figsize=(14, 10))
            axes = axes.flatten()

            for i, col in enumerate(numeric_cols[:4]):
                axes[i].hist(df[col].dropna(), bins=30, edgecolor='black', alpha=0.7)
                axes[i].set_title(f'Distribution: {col}', fontweight='bold')
                axes[i].set_xlabel(col)
                axes[i].set_ylabel('Frequency')

            for i in range(len(numeric_cols[:4]), 4):
                axes[i].axis('off')

            plt.tight_layout()
            path = self.charts_dir / 'distributions.png'
            safe_savefig(fig, path, desired_dpi=150, min_dpi=72)
            plt.close(fig)
            chart_paths['distributions'] = path

        if results and 'feature_importance' in results:
            importance = results['feature_importance']
            if importance:
                features = list(importance.keys())[:15]
                values = [importance[f] for f in features]

                fig, ax = plt.subplots(figsize=(10, 8))
                ax.barh(features, values, color='steelblue')
                ax.set_xlabel('Importance', fontweight='bold')
                ax.set_title('Top 15 Feature Importances', fontsize=14, fontweight='bold')
                ax.invert_yaxis()
                plt.tight_layout()
                path = self.charts_dir / 'feature_importance.png'
                safe_savefig(fig, path, desired_dpi=150, min_dpi=72)
                plt.close(fig)
                chart_paths['feature_importance'] = path

        if results and 'models' in results:
            model_names = list(results['models'].keys())
            if results.get('task_type') == 'classification':
                scores = [results['models'][m].get('accuracy', 0) for m in model_names]
                metric = 'Accuracy'
            else:
                scores = [results['models'][m].get('r2_score', 0) for m in model_names]
                metric = 'R² Score'

            fig, ax = plt.subplots(figsize=(8, 5))
            colors_list = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728']
            bars = ax.bar(model_names, scores, color=[colors_list[i % len(colors_list)] for i in range(len(model_names))])
            ax.set_ylabel(metric, fontweight='bold')
            ax.set_title(f'Model Performance Comparison ({metric})', fontsize=14, fontweight='bold')
            max_score = max(scores) if scores else 1.0
            ax.set_ylim([0, max_score * 1.1])

            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width() / 2., height,
                        f'{height:.3f}', ha='center', va='bottom', fontweight='bold')

            plt.xticks(rotation=45, ha='right')
            plt.tight_layout()
            path = self.charts_dir / 'model_comparison.png'
            safe_savefig(fig, path, desired_dpi=150, min_dpi=72)
            plt.close(fig)
            chart_paths['model_comparison'] = path

        return chart_paths

    def _generate_pdf_report(self, df: pd.DataFrame, metadata: Dict[str, Any],
                             results: Dict[str, Any], chart_paths: Dict[str, Path],
                             output_path: Path):

        doc = SimpleDocTemplate(str(output_path), pagesize=letter,
                               rightMargin=72, leftMargin=72,
                               topMargin=72, bottomMargin=18)

        story = []
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=30,
            alignment=1
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=12,
            spaceBefore=12
        )

        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            leading=16,
            spaceAfter=12
        )

        story.append(Paragraph("DataPrepX Analysis Report", title_style))
        story.append(Spacer(1, 12))
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(Spacer(1, 30))

        story.append(Paragraph("1. Executive Summary", heading_style))

        if results and 'ai_summary' in results:
            summary_text = results['ai_summary']
            summary_text = summary_text.replace('###', '').replace('##', '').replace('**', '')
            summary_text = summary_text.replace('■', '').replace('▪', '').replace('•', '')
            summary_text = summary_text.replace('\u2011', '-').replace('\xd7', 'x')
            summary_text = summary_text.replace('\u2192', '->').replace('\xb2', '²')
            
            summary_paragraphs = summary_text.split('\n\n')
            for para in summary_paragraphs:
                if para.strip():
                    clean_para = para.strip().replace('\n', ' ')
                    clean_para = ' '.join(clean_para.split())
                    if clean_para and not clean_para.startswith(('#', '-', '*', '.')):
                        story.append(Paragraph(clean_para, body_style))
        else:
            summary_text = f"""
            This report provides a comprehensive analysis of the dataset processed through DataPrepX.
            The dataset contains {metadata['final_shape'][0]:,} rows and {metadata['final_shape'][1]} features
            after preprocessing and feature engineering.
            """
            story.append(Paragraph(summary_text, body_style))

        story.append(Spacer(1, 20))

        story.append(Paragraph("2. Data Overview", heading_style))

        data_table_data = [
            ['Metric', 'Value'],
            ['Original Rows', f"{metadata['original_shape'][0]:,}"],
            ['Original Columns', str(metadata['original_shape'][1])],
            ['Final Rows', f"{metadata['final_shape'][0]:,}"],
            ['Final Columns', str(metadata['final_shape'][1])],
            ['Duplicates Removed', str(metadata.get('duplicates_removed', 0))]
        ]

        data_table = Table(data_table_data, colWidths=[3 * inch, 3 * inch])
        data_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 13),
            ('FONTSIZE', (0, 1), (-1, -1), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 14),
            ('TOPPADDING', (0, 1), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 10),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8f9ff')),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e0e0e0')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
        ]))

        story.append(data_table)
        story.append(Spacer(1, 20))

        if metadata.get('missing_values'):
            story.append(Paragraph("3. Missing Values Handled", heading_style))
            
            missing_items = list(metadata['missing_values'].items())[:10]
            if missing_items:
                missing_table_data = [['Column', 'Missing Count']]
                for col, count in missing_items:
                    missing_table_data.append([str(col), str(count)])
                
                missing_table = Table(missing_table_data, colWidths=[3.5 * inch, 2.5 * inch])
                missing_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('FONTSIZE', (0, 1), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('TOPPADDING', (0, 1), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8f9ff')),
                    ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e0e0e0')),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
                ]))
                story.append(missing_table)
            story.append(Spacer(1, 20))

        story.append(PageBreak())

        story.append(Paragraph("4. Data Visualizations", heading_style))

        if 'distributions' in chart_paths:
            story.append(Paragraph("4.1 Feature Distributions", styles['Heading3']))
            story.append(safe_load_image(chart_paths['distributions'], width=5.5, height=3.5))
            story.append(Spacer(1, 20))

        if 'correlation' in chart_paths:
            story.append(Paragraph("4.2 Correlation Analysis", styles['Heading3']))
            story.append(safe_load_image(chart_paths['correlation'], width=5.5, height=4))
            story.append(Spacer(1, 20))

        story.append(PageBreak())

        if results:
            story.append(Paragraph("5. Machine Learning Results", heading_style))

            ml_summary_data = [
                ['Metric', 'Value'],
                ['Task Type', results['task_type'].title()],
                ['Target Column', results['target_column']],
                ['Training Samples', f"{results['train_size']:,}"],
                ['Test Samples', f"{results['test_size']:,}"],
                ['Best Model', results['best_model']],
                ['Best Score', f"{results['best_score']:.4f}"]
            ]
            
            ml_summary_table = Table(ml_summary_data, colWidths=[2.5 * inch, 3.5 * inch])
            ml_summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('TOPPADDING', (0, 1), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8f9ff')),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e0e0e0')),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
            ]))
            story.append(ml_summary_table)
            story.append(Spacer(1, 20))

            if results.get('models'):
                story.append(Paragraph("5.1 Detailed Model Metrics", styles['Heading3']))

                if results['task_type'] == 'classification':
                    metrics_data = [['Model', 'Accuracy', 'Precision', 'Recall', 'F1-Score', 'CV Mean']]
                    for model_name, metrics in results['models'].items():
                        metrics_data.append([
                            model_name,
                            f"{metrics.get('accuracy', 0):.4f}",
                            f"{metrics.get('precision', 0):.4f}",
                            f"{metrics.get('recall', 0):.4f}",
                            f"{metrics.get('f1_score', 0):.4f}",
                            f"{metrics.get('cv_mean', 0):.4f}"
                        ])
                else:
                    metrics_data = [['Model', 'R² Score', 'RMSE', 'MAE', 'CV Mean']]
                    for model_name, metrics in results['models'].items():
                        metrics_data.append([
                            model_name,
                            f"{metrics.get('r2_score', 0):.4f}",
                            f"{metrics.get('rmse', 0):.2f}",
                            f"{metrics.get('mae', 0):.2f}",
                            f"{metrics.get('cv_mean', 0):.4f}"
                        ])

                col_width = 6.5 * inch / len(metrics_data[0])
                metrics_table = Table(metrics_data, colWidths=[col_width] * len(metrics_data[0]))
                metrics_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('TOPPADDING', (0, 1), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8f9ff')),
                    ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e0e0e0')),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
                ]))

                story.append(metrics_table)
                story.append(Spacer(1, 20))

            if 'model_comparison' in chart_paths:
                story.append(Paragraph("5.2 Model Performance", styles['Heading3']))
                story.append(safe_load_image(chart_paths['model_comparison'], width=5.5, height=3))
                story.append(Spacer(1, 20))

            if 'feature_importance' in chart_paths:
                story.append(Paragraph("5.3 Feature Importance", styles['Heading3']))
                story.append(safe_load_image(chart_paths['feature_importance'], width=5.5, height=4))
                story.append(Spacer(1, 20))

        story.append(PageBreak())
        story.append(Paragraph("6. Conclusion", heading_style))

        if results:
            conclusion = f"""
            The DataPrepX pipeline successfully processed the dataset, applying comprehensive
            preprocessing, feature engineering, and machine learning modeling.
            The best performing model was {results['best_model']} with a score of {results['best_score']:.4f}.
            All generated visualizations and detailed metrics are included in this report.
            """
        else:
            conclusion = """
            The DataPrepX pipeline successfully processed the dataset, applying comprehensive
            preprocessing and feature engineering. All generated visualizations and detailed 
            metrics are included in this report.
            """

        story.append(Paragraph(conclusion, body_style))

        doc.build(story)

    def _generate_docx_report(self, df: pd.DataFrame, metadata: Dict[str, Any],
                              results: Dict[str, Any], chart_paths: Dict[str, Path],
                              output_path: Path):

        doc = Document()

        title = doc.add_heading('DataPrepX Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()

        doc.add_heading('1. Executive Summary', 1)

        if results and 'ai_summary' in results:
            summary_text = results['ai_summary']
            summary_text = summary_text.replace('###', '').replace('##', '').replace('**', '')
            summary_text = summary_text.replace('■', '').replace('▪', '').replace('•', '')
            summary_text = summary_text.replace('\u2011', '-').replace('\xd7', 'x')
            summary_text = summary_text.replace('\u2192', '->').replace('\xb2', '²')
            
            summary_paragraphs = summary_text.split('\n\n')
            for para in summary_paragraphs:
                if para.strip():
                    clean_text = para.strip()
                    clean_text = ' '.join(clean_text.split())
                    if clean_text and not clean_text.startswith(('#', '-', '*', '.')):
                        doc.add_paragraph(clean_text)
        else:
            doc.add_paragraph(
                f"This report provides a comprehensive analysis of the dataset processed through DataPrepX. "
                f"The dataset contains {metadata['final_shape'][0]:,} rows and {metadata['final_shape'][1]} features "
                f"after preprocessing and feature engineering."
            )

        doc.add_heading('2. Data Overview', 1)
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Metric'
        table.rows[0].cells[1].text = 'Value'
        table.rows[1].cells[0].text = 'Original Rows'
        table.rows[1].cells[1].text = f"{metadata['original_shape'][0]:,}"
        table.rows[2].cells[0].text = 'Original Columns'
        table.rows[2].cells[1].text = str(metadata['original_shape'][1])
        table.rows[3].cells[0].text = 'Final Rows'
        table.rows[3].cells[1].text = f"{metadata['final_shape'][0]:,}"
        table.rows[4].cells[0].text = 'Final Columns'
        table.rows[4].cells[1].text = str(metadata['final_shape'][1])
        table.rows[5].cells[0].text = 'Duplicates Removed'
        table.rows[5].cells[1].text = str(metadata.get('duplicates_removed', 0))

        doc.add_paragraph()

        if metadata.get('missing_values'):
            doc.add_heading('3. Missing Values Handled', 1)
            doc.add_paragraph("The following columns had missing values that were imputed:")
            for col, count in list(metadata['missing_values'].items())[:10]:
                doc.add_paragraph(f"{col}: {count} missing values", style='List Bullet')

        doc.add_page_break()

        doc.add_heading('4. Data Visualizations', 1)

        if 'distributions' in chart_paths:
            doc.add_heading('4.1 Feature Distributions', 2)
            doc.add_picture(str(chart_paths['distributions']), width=Inches(6))
            doc.add_paragraph()

        if 'correlation' in chart_paths:
            doc.add_heading('4.2 Correlation Analysis', 2)
            doc.add_picture(str(chart_paths['correlation']), width=Inches(6))
            doc.add_paragraph()

        doc.add_page_break()

        if results:
            doc.add_heading('5. Machine Learning Results', 1)

            doc.add_paragraph(f"Task Type: {results['task_type'].title()}")
            doc.add_paragraph(f"Target Column: {results['target_column']}")
            doc.add_paragraph(f"Training Samples: {results['train_size']:,}")
            doc.add_paragraph(f"Test Samples: {results['test_size']:,}")
            doc.add_paragraph(f"Best Model: {results['best_model']}")
            doc.add_paragraph(f"Performance Score: {results['best_score']:.4f}")
            doc.add_paragraph()

            if results.get('models'):
                doc.add_heading('5.1 Detailed Model Metrics', 2)

                if results['task_type'] == 'classification':
                    table = doc.add_table(rows=len(results['models']) + 1, cols=6)
                    table.style = 'Light Grid Accent 1'

                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Model'
                    hdr_cells[1].text = 'Accuracy'
                    hdr_cells[2].text = 'Precision'
                    hdr_cells[3].text = 'Recall'
                    hdr_cells[4].text = 'F1-Score'
                    hdr_cells[5].text = 'CV Mean'

                    for idx, (model_name, metrics) in enumerate(results['models'].items(), 1):
                        row_cells = table.rows[idx].cells
                        row_cells[0].text = model_name
                        row_cells[1].text = f"{metrics.get('accuracy', 0):.4f}"
                        row_cells[2].text = f"{metrics.get('precision', 0):.4f}"
                        row_cells[3].text = f"{metrics.get('recall', 0):.4f}"
                        row_cells[4].text = f"{metrics.get('f1_score', 0):.4f}"
                        row_cells[5].text = f"{metrics.get('cv_mean', 0):.4f}"
                else:
                    table = doc.add_table(rows=len(results['models']) + 1, cols=5)
                    table.style = 'Light Grid Accent 1'

                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Model'
                    hdr_cells[1].text = 'R² Score'
                    hdr_cells[2].text = 'RMSE'
                    hdr_cells[3].text = 'MAE'
                    hdr_cells[4].text = 'CV Mean'

                    for idx, (model_name, metrics) in enumerate(results['models'].items(), 1):
                        row_cells = table.rows[idx].cells
                        row_cells[0].text = model_name
                        row_cells[1].text = f"{metrics.get('r2_score', 0):.4f}"
                        row_cells[2].text = f"{metrics.get('rmse', 0):.2f}"
                        row_cells[3].text = f"{metrics.get('mae', 0):.2f}"
                        row_cells[4].text = f"{metrics.get('cv_mean', 0):.4f}"

                doc.add_paragraph()

            if 'model_comparison' in chart_paths:
                doc.add_heading('5.2 Model Performance', 2)
                doc.add_picture(str(chart_paths['model_comparison']), width=Inches(6))
                doc.add_paragraph()

            if 'feature_importance' in chart_paths:
                doc.add_heading('5.3 Feature Importance', 2)
                doc.add_picture(str(chart_paths['feature_importance']), width=Inches(6))
                doc.add_paragraph()

        doc.add_page_break()

        doc.add_heading('6. Conclusion', 1)

        if results:
            conclusion_text = (
                "The DataPrepX pipeline successfully processed the dataset, applying comprehensive "
                "preprocessing, feature engineering, and machine learning modeling. "
                f"The best performing model was {results['best_model']} with a score of {results['best_score']:.4f}. "
                "All generated visualizations and detailed metrics are included in this report."
            )
        else:
            conclusion_text = (
                "The DataPrepX pipeline successfully processed the dataset, applying comprehensive "
                "preprocessing and feature engineering. All generated visualizations and detailed "
                "metrics are included in this report."
            )

        doc.add_paragraph(conclusion_text)

        doc.save(str(output_path))